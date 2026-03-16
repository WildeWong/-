from .base import BaseParser, ParseResult


class DocxParser(BaseParser):
    """Microsoft Word (.docx) script parser with style-aware scene detection."""

    # Word style names that directly indicate a scene heading
    SCENE_HEADING_STYLES = frozenset({
        "scene heading", "scene heading char",
        "shot", "scene",
        "heading 1", "heading 2",           # some templates use headings for scenes
        "标题 1", "标题1", "场景标题",       # Chinese Word style names
    })

    def parse(self, file_path: str) -> ParseResult:
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for .docx parsing. Install with: pip install python-docx"
            )

        doc = Document(file_path)
        lines: list[str] = []
        line_metadata: dict[int, dict[str, str]] = {}

        def _process_para(para) -> None:
            text = para.text  # preserve internal spacing
            style_name = (para.style.name or "").strip() if para.style else ""
            line_idx = len(lines)
            lines.append(text)

            # Mark as confirmed Scene Heading when style name matches
            if style_name.lower() in self.SCENE_HEADING_STYLES:
                line_metadata[line_idx] = {
                    "type": "Scene Heading",
                    "style": style_name,
                }

        # Main body paragraphs
        for para in doc.paragraphs:
            _process_para(para)

        # Tables — iterate each cell's paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            _process_para(para)

        # Clean up: strip trailing whitespace from each line
        lines = [ln.rstrip() for ln in lines]

        # Collapse long runs of blank lines to at most one
        cleaned: list[str] = []
        blank_run = 0
        for ln in lines:
            if not ln.strip():
                blank_run += 1
                if blank_run > 1:
                    continue
            else:
                blank_run = 0
            cleaned.append(ln)

        # Remap line_metadata indices after blank-line collapsing
        # Build old->new index map from the cleaning step
        new_meta: dict[int, dict[str, str]] = {}
        new_idx = 0
        blank_run2 = 0
        for old_idx, ln in enumerate(lines):
            if not ln.strip():
                blank_run2 += 1
                if blank_run2 > 1:
                    continue
            else:
                blank_run2 = 0
            if old_idx in line_metadata:
                new_meta[new_idx] = line_metadata[old_idx]
            new_idx += 1

        return ParseResult(lines=cleaned, line_metadata=new_meta)

    @staticmethod
    def supported_extensions() -> list[str]:
        return [".docx"]
